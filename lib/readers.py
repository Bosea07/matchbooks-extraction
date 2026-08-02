import io, csv, re

def read_csv(data: bytes):
    text = None
    for enc in ('utf-8-sig', 'utf-8', 'cp1256', 'latin-1'):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise ValueError('Could not decode CSV file')
    sample = text[:4000]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')
    except csv.Error:
        class dialect:
            delimiter = ','
            quotechar = '"'
    rows = [r for r in csv.reader(io.StringIO(text), delimiter=dialect.delimiter, quotechar=getattr(dialect, 'quotechar', '"') or '"')]
    return rows, {'reader': 'csv'}

def _score_sheet(grid):
    from .normalize import parse_amount
    score = 0
    for row in grid[:60]:
        nums = sum(1 for c in row if parse_amount(c) is not None)
        if nums >= 1 and sum(1 for c in row if c not in (None, '')) >= 2:
            score += 1 + min(nums, 3)
    return score

def read_xlsx(data: bytes):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    best, best_grid, best_score = None, [], -1
    for ws in wb.worksheets:
        grid = [[c for c in row] for row in ws.iter_rows(values_only=True, max_row=5000)]
        s = _score_sheet(grid)
        if s > best_score:
            best, best_grid, best_score = ws.title, grid, s
    return best_grid, {'reader': 'xlsx', 'sheet': best}

def read_xls(data: bytes):
    import xlrd
    wb = xlrd.open_workbook(file_contents=data)
    best, best_grid, best_score = None, [], -1
    for ws in wb.sheets():
        grid = []
        for r in range(min(ws.nrows, 5000)):
            row = []
            for c in range(ws.ncols):
                cell = ws.cell(r, c)
                if cell.ctype == 3:
                    try:
                        import xlrd.xldate as xd
                        row.append(xd.xldate_as_datetime(cell.value, wb.datemode))
                    except Exception:
                        row.append(cell.value)
                else:
                    row.append(cell.value if cell.value != '' else None)
            grid.append(row)
        s = _score_sheet(grid)
        if s > best_score:
            best, best_grid, best_score = ws.name, grid, s
    return best_grid, {'reader': 'xls', 'sheet': best}

def read_docx(data: bytes):
    import docx
    doc = docx.Document(io.BytesIO(data))
    grid = []
    for table in doc.tables:
        for row in table.rows:
            grid.append([cell.text.strip() or None for cell in row.cells])
    if not grid:
        for p in doc.paragraphs:
            line = p.text.strip()
            if line:
                grid.append(_split_line(line))
    return grid, {'reader': 'docx', 'tables': len(doc.tables)}

_MULTISPACE = re.compile(r'\s{2,}|\t')

def _words_to_rows(words, line_tol=3.0):
    lines = []
    for w in sorted(words, key=lambda w: (round(w['top'], 1), w['x0'])):
        for ln in lines:
            if abs(ln['top'] - w['top']) <= line_tol:
                ln['words'].append(w)
                break
        else:
            lines.append({'top': w['top'], 'words': [w]})
    rows = []
    for ln in lines:
        ws = sorted(ln['words'], key=lambda w: w['x0'])
        widths = [(w['x1'] - w['x0']) / max(len(w['text']), 1) for w in ws]
        cw = sorted(widths)[len(widths) // 2] if widths else 5.0
        gap_thresh = max(cw * 1.8, 7.0)
        cells, cur = [], ws[0]['text']
        for a, b in zip(ws, ws[1:]):
            if b['x0'] - a['x1'] > gap_thresh:
                cells.append(cur)
                cur = b['text']
            else:
                cur += ' ' + b['text']
        cells.append(cur)
        rows.append([c.strip() or None for c in cells])
    return rows

def _split_line(line):
    parts = _MULTISPACE.split(line.strip())
    if len(parts) < 2:
        parts = line.strip().split()
    return [p.strip() or None for p in parts]

def read_pdf(data: bytes):
    import pdfplumber
    grid, used_tables, text_pages = [], 0, 0
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            good = [t for t in tables if t and len(t) >= 2]
            if good:
                used_tables += len(good)
                for t in good:
                    for row in t:
                        grid.append([(c.strip() if isinstance(c, str) else c) or None for c in row])
            else:
                words = page.extract_words() or []
                if words:
                    text_pages += 1
                    grid.extend(_words_to_rows(words))
    scanned = not grid
    return grid, {'reader': 'pdf', 'tables': used_tables, 'textPages': text_pages, 'scanned': scanned}

READERS = {
    'csv': read_csv, 'txt': read_csv,
    'xlsx': read_xlsx, 'xlsm': read_xlsx, 'xltx': read_xlsx,
    'xls': read_xls,
    'docx': read_docx,
    'pdf': read_pdf,
}

def read_any(filename: str, data: bytes):
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext == 'doc':
        raise ValueError("Legacy .doc files aren't supported — please save the statement as .docx or PDF and re-upload.")
    if ext not in READERS:
        raise ValueError(f"Unsupported file type '.{ext}'. Supported: CSV, XLSX, XLS, PDF, DOCX.")
    return READERS[ext](data)
